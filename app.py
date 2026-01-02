import streamlit as st
import pdfplumber
from pdf2docx import Converter
from docx import Document
import io
import pytesseract
from PIL import Image

st.set_page_config(page_title="PDF to Word", layout="wide", initial_sidebar_state="expanded")
st.title("🔥 PDF → Word Converter **GRATIS**")
st.markdown("### Sube tu PDF → Preview → Descarga Word editable")

# Sidebar
st.sidebar.header("📱 Móvil OK")
st.sidebar.markdown("""
- ✅ iPhone/Android
- ✅ Preview página 1  
- ✅ OCR automático
- ✅ Todas las páginas
""")

# Upload
uploaded_file = st.file_uploader("📤 **Arrastra tu PDF aquí**", type=["pdf"])

if uploaded_file:
    # Preview primera página
    with pdfplumber.open(uploaded_file) as pdf:
        if pdf.pages:
            st.image(pdf.pages[0].to_image(resolution=120).original, 
                    caption="**Preview página 1**", use_column_width=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("**🚀 Convertir NORMAL** (formato preservado)", type="primary"):
            with st.spinner("Convirtiendo... ⏳"):
                cv = Converter(uploaded_file)
                cv.convert("output.docx")
                cv.close()
                
                with open("output.docx", "rb") as f:
                    st.download_button(
                        label="📥 **DESCARGAR WORD**",
                        data=f.read(),
                        file_name="pdf_convertido.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
    
    with col2:
        if st.button("**🖼️ OCR ESCANEADO** (imágenes a texto)", type="secondary"):
            with st.spinner("Procesando OCR... 🤖"):
                doc = Document()
                with pdfplumber.open(uploaded_file) as pdf:
                    for i, page in enumerate(pdf.pages[:5]):  # Primeras 5 páginas demo
                        image = page.to_image(resolution=200).original
                        text = pytesseract.image_to_string(image, lang="spa")
                        doc.add_paragraph(f"--- PÁGINA {i+1} ---")
                        doc.add_paragraph(text)
                
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.download_button(
                    label="📥 **DESCARGAR WORD OCR**",
                    data=buffer,
                    file_name="pdf_ocr.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
    st.balloons()
    st.success("✅ **¡Listo! Usa en móvil añadiendo a pantalla inicio**")
